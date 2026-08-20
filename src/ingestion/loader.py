import fitz
import os
import re
from typing import List, Dict, Optional
from docx import Document
import openpyxl

# Détecte des codes de référence type SMQ-FOR-092-A, ISO-9001, PROC-QUAL-01, etc.
# Générique : pas spécifique à un format d'entreprise particulier.
REFERENCE_PATTERN = re.compile(r"\b[A-Z]{2,6}(?:-[A-Z0-9]{2,6}){1,3}\b")

# Gain rapide #1 (EF-DOC-02) : détection heuristique du statut du document
# à partir de mots-clés dans l'en-tête/pied de page. Approximatif par nature
# (pas de champ structuré dédié côté source) — affiné si un jour la GED
# expose un vrai statut en métadonnée plutôt qu'en texte libre.
STATUS_PATTERNS = {
    "obsolete": re.compile(r"obsol[eè]te|annul[eé]|p[eé]rim[eé]|remplac[eé]\s+par", re.IGNORECASE),
    "en_revision": re.compile(r"en\s+r[eé]vision|brouillon|draft|en\s+cours\s+de\s+r[eé]daction|projet\s+de\s+version", re.IGNORECASE),
}


def detect_status(text: str) -> str:
    if not text:
        return "en_vigueur"
    if STATUS_PATTERNS["obsolete"].search(text):
        return "obsolete"
    if STATUS_PATTERNS["en_revision"].search(text):
        return "en_revision"
    return "en_vigueur"


def clean_title(file_path: str) -> str:
    name = os.path.splitext(os.path.basename(file_path))[0]
    name = name.replace("_", " ")
    return " ".join(name.split())


def extract_header_footer(doc: Document) -> str:
    """python-docx ignore les en-têtes/pieds de page via .paragraphs/.tables -
    on les récupère explicitement ici, car c'est souvent là que vivent
    les références et versions des documents qualité."""
    parts = []
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                t = p.text.strip()
                if t:
                    parts.append(t)
            for table in container.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
    # dédoublonnage en gardant l'ordre (Word répète souvent le header par section)
    seen = set()
    unique_parts = []
    for p in parts:
        if p not in seen:
            seen.add(p)
            unique_parts.append(p)
    return "\n".join(unique_parts)


def detect_reference(text: str) -> Optional[str]:
    match = REFERENCE_PATTERN.search(text)
    return match.group(0) if match else None


def load_pdf(file_path: str) -> List[Dict]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")
    documents = []
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()
        if text:
            # Statut par défaut "en_vigueur" : pas encore de détection pour les PDF
            # (contrairement au docx, l'en-tête/pied de page structuré n'est pas
            # exploité de la même façon avec PyMuPDF — amélioration possible plus tard).
            documents.append({"text": text, "metadata": {"source": os.path.basename(file_path), "page": page_num + 1, "total_pages": len(doc), "status": "en_vigueur"}})
    doc.close()
    return documents


def load_docx(file_path: str) -> List[Dict]:
    doc = Document(file_path)
    documents = []
    source = os.path.basename(file_path)
    title = clean_title(file_path)

    header_footer_text = extract_header_footer(doc)
    doc_ref = detect_reference(header_footer_text) if header_footer_text else None
    prefix = f"[{doc_ref}] " if doc_ref else ""
    doc_ref_meta = doc_ref or ""  # ChromaDB rejette None dans les métadonnées
    doc_status = detect_status(header_footer_text)

    # Chunk dédié pour l'en-tête/pied de page : c'est souvent la SEULE
    # source de la référence et de la version du document.
    if header_footer_text:
        documents.append({
            "text": f"{title} — En-tête du document : {header_footer_text}",
            "metadata": {"source": source, "page": 1, "total_pages": 1, "type": "header_footer", "doc_ref": doc_ref_meta, "status": doc_status}
        })

    paragraph_blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraph_blocks:
        documents.append({
            "text": prefix + "\n".join(paragraph_blocks),
            "metadata": {"source": source, "page": 1, "total_pages": 1, "doc_ref": doc_ref_meta, "status": doc_status}
        })

    for table_num, table in enumerate(doc.tables, start=1):
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        for row_num, row in enumerate(table.rows[1:] if len(table.rows) > 1 else table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            pairs = []
            for header, cell in zip(headers, cells):
                if cell:
                    # Si l'en-tête de colonne EST le titre du document (arrive quand
                    # une table "identité" utilise le nom du doc comme libellé de
                    # colonne au lieu d'un label générique), ne pas le répéter : le
                    # titre est déjà dans le préfixe de chaque ligne. Sans ce garde-fou,
                    # certaines lignes répètent le titre 2-3 fois, ce qui les fait
                    # ressembler à un contenu très pertinent pour n'importe quelle
                    # question sur ce document, même quand leur contenu réel ne l'est
                    # pas — confirmé sur RAGAS (context_precision) le 2026-08.
                    if header and header.strip().lower() == title.strip().lower():
                        pairs.append(cell)
                    else:
                        pairs.append(f"{header}: {cell}" if header else cell)
            if not pairs:
                pairs = [c for c in cells if c]
            if pairs:
                documents.append({
                    "text": f"{prefix}{title} — " + " | ".join(pairs),
                    "metadata": {"source": source, "page": 1, "total_pages": 1, "table": table_num, "row": row_num, "doc_ref": doc_ref_meta, "status": doc_status}
                })

    return documents


def load_xlsx(file_path: str) -> List[Dict]:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    documents = []
    title = clean_title(file_path)
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h).strip() if h is not None else "" for h in rows[0]]
        for row_num, row in enumerate(rows[1:], start=2):
            pairs = []
            for header, cell in zip(headers, row):
                if cell is not None and str(cell).strip():
                    pairs.append(f"{header}: {cell}" if header else str(cell))
            if pairs:
                row_text = f"{title} ({sheet}) — " + " | ".join(pairs)
                documents.append({
                    "text": row_text,
                    "metadata": {"source": os.path.basename(file_path), "page": 1, "total_pages": 1, "sheet": sheet, "row": row_num, "status": "en_vigueur"}
                })
    return documents


def load_file(file_path: str) -> List[Dict]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".xlsx":
        return load_xlsx(file_path)
    return []


def load_folder(folder_path: str) -> List[Dict]:
    all_docs = []
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.endswith((".pdf", ".docx", ".xlsx")):
                file_path = os.path.join(root, filename)
                rel_path = os.path.relpath(file_path, folder_path).replace(os.sep, "/")
                print(f"Chargement : {filename}")
                docs = load_file(file_path)
                for d in docs:
                    d["metadata"]["path"] = rel_path
                all_docs.extend(docs)
    return all_docs